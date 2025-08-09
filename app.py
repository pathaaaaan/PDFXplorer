import os
import base64
from flask import Flask, request, jsonify, send_from_directory, render_template, send_file, Response
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import PyPDF2
import io
import tempfile
import re
import csv
from pdf2image import convert_from_path
import pytesseract
import pdfplumber
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert as docx2pdf_convert

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB
ALLOWED_EXTENSIONS = {'pdf'}

# Load weak-strong word dictionary
word_dict = {}
with open('weak_strong_words.csv', 'r', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        word_dict[row['weak_word'].lower()] = row['strong_word']

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/list', methods=['GET'])
def list_files():
    files = []
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        if filename.endswith('.pdf'):
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            files.append({
                'name': filename,
                'size': os.path.getsize(path),
                'modified': os.path.getmtime(path)
            })
    return jsonify(files)

@app.route('/pdf/<filename>')
def serve_pdf(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/pdf/<filename>/metadata')
def get_pdf_metadata(filename):
    try:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        doc = fitz.open(filepath)
        metadata = {
            'total_pages': len(doc),
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', '')
        }
        doc.close()
        return jsonify(metadata)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete/<filename>', methods=['DELETE'])
def delete_file(filename):
    try:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
        if os.path.exists(filepath):
            os.remove(filepath)
            return jsonify({'message': 'File deleted successfully'})
        return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/replace', methods=['POST'])
def replace_text():
    data = request.json
    filename = data.get('filename')
    search_text = data.get('searchText')
    replace_text = data.get('replaceText')
    page_num = data.get('pageNum')
    match_index = data.get('matchIndex')
    ignore_case = data.get('ignoreCase', True)
    
    if not all([filename, search_text, replace_text is not None, page_num is not None, match_index is not None]):
        return jsonify({'error': 'Missing parameters'}), 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    try:
        def normalize_text(text):
            # Replace non-breaking spaces and other common whitespace variants
            text = text.replace('\u00A0', ' ').replace('\r\n', ' ').replace('\n', ' ')
            # Normalize unicode characters
            text = text.encode('unicode_escape').decode('ascii', 'ignore')
            return text.strip()

        # Create a temporary file for the modified PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_path = temp_file.name

        # Open the source PDF
        with fitz.open(filepath) as doc:
            if page_num > len(doc) or page_num < 1:
                return jsonify({'error': 'Page number out of range'}), 400

            page = doc[page_num - 1]
            
            # Get page text with layout preservation
            text = page.get_text('text')
            
            # Prepare search pattern with word boundaries and optional case sensitivity
            flags = re.IGNORECASE if ignore_case else 0
            pattern = re.escape(search_text)
            matches = list(re.finditer(pattern, text, flags))
            
            if not matches:
                # Try with normalized text if no matches found
                normalized_text = normalize_text(text)
                normalized_search = normalize_text(search_text)
                matches = list(re.finditer(re.escape(normalized_search), normalized_text, flags))
                text = normalized_text
                
            if match_index >= len(matches):
                return jsonify({
                    'error': 'Match not found', 
                    'debug': {
                        'match_count': len(matches),
                        'page_text_preview': text[:200] + '...' if len(text) > 200 else text
                    }
                }), 404

            # Perform the replacement
            match = matches[match_index]
            new_text = text[:match.start()] + replace_text + text[match.end():]
            
           
            # we'll just save the modified text as an annotation
            rect = fitz.Rect(72, 72, 200, 100)  
            annot = page.add_freetext_annot(rect, new_text)
            annot.update()
            
            # Save the modified document
            doc.save(temp_path)
        
        # Replace the original file with the modified one
        os.replace(temp_path, filepath)
        
        return jsonify({
            'message': 'Text replaced successfully',
            'details': {
                'original_text': match.group(),
                'replaced_with': replace_text,
                'position': {
                    'start': match.start(),
                    'end': match.end()
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            'error': 'Failed to replace text',
            'details': str(e),
            'type': type(e).__name__
        }), 500

@app.route('/download-annotated', methods=['POST'])
def download_annotated():
    try:
        data = request.json
        filename = data.get('filename')
        annotations = data.get('annotations', [])
        if not filename:
            return jsonify({'error': 'Filename required'}), 400
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
            
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_path = temp_file.name
            
        doc = fitz.open(filepath)
        for page_num, page_annots in annotations:
            if page_num <= len(doc):
                page = doc[page_num - 1]
                for ann in page_annots:
                    rect = fitz.Rect(ann['x'], ann['y'], ann['x'] + ann['width'], ann['y'] + ann['height'])
                    if ann['type'] == 'highlight':
                        highlight = page.add_highlight_annot(rect)
                        highlight.set_colors(stroke=ann['color'])
                        highlight.update()
                    elif ann['type'] == 'text-note':
                        note = page.add_text_annot(rect.tl, ann['text'])
                        note.set_colors(stroke=ann['color'])
                        note.update()
                    elif ann['type'] in ['rectangle', 'circle']:
                        shape = page.add_rect_annot(rect) if ann['type'] == 'rectangle' else page.add_circle_annot(rect)
                        shape.set_colors(stroke=ann['color'])
                        shape.update()
        doc.save(temp_path)
        doc.close()
        return send_file(temp_path, as_attachment=True, download_name=f"annotated_{filename}", mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if 'temp_path' in locals():
            try: os.unlink(temp_path)
            except: pass

def create_temp_copy(src_path):
    """Create a temporary copy of a file that might be locked by another process."""
    import shutil
    import tempfile
    
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, os.path.basename(src_path))
    
    try:
        # Try to copy using shutil first
        shutil.copy2(src_path, temp_path)
        return temp_path, temp_dir
    except Exception as e:
        # If that fails, try reading in binary mode and writing to temp file
        try:
            with open(src_path, 'rb') as src_file:
                content = src_file.read()
            with open(temp_path, 'wb') as dst_file:
                dst_file.write(content)
            return temp_path, temp_dir
        except Exception as e:
            # Clean up temp directory if we created it
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            raise e

@app.route('/analyze-words', methods=['POST'])
def analyze_words():
    temp_path = None
    temp_dir = None
    try:
        data = request.json
        filename = data.get('filename')
        page_num = data.get('pageNum')
        action = data.get('action', 'highlight')
        
        if not filename or not page_num:
            return jsonify({'error': 'Missing filename or pageNum'}), 400
            
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        # Create a temporary copy of the file to work with
        temp_filepath, temp_dir = create_temp_copy(filepath)
        
        try:
            with fitz.open(temp_filepath) as doc:  # Use context manager for automatic cleanup
                if page_num > len(doc) or page_num < 1:
                    return jsonify({'error': 'Page number out of range'}), 400
                    
                page = doc[page_num - 1]
                text = page.get_text()
                weak_words = []

                for weak_word in word_dict.keys():
                    for match in re.finditer(r'\b' + re.escape(weak_word) + r'\b', text.lower()):
                        weak_words.append({
                            'word': match.group(),
                            'strong_word': word_dict[weak_word],
                            'start': match.start(),
                            'end': match.end()
                        })
                
                if action == 'replace':
                    # Create a temporary file for the modified version
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                        temp_path = temp_file.name
                    
                    # Create a copy of the document to modify
                    new_doc = fitz.open()
                    new_doc.insert_pdf(doc, from_page=page_num-1, to_page=page_num-1)
                    new_page = new_doc[0]
                    
                    # Get the default font for consistent text rendering
                    font_size = 11
                    
                    for ww in weak_words:
                        # Find all instances of the weak word
                        rects = new_page.search_for(ww['word'])
                        
                        # Process each found instance
                        for rect in rects:
                            # Get the bounding box of the found text
                            bbox = rect
                            
                            # Calculate text width to determine if we need to adjust font size
                            strong_text = ww['strong_word']
                            
                            # Get the original text's font properties if possible
                            text_instances = new_page.search_for(ww['word'])
                            if text_instances:
                                # Try to get font size from the original text
                                text_instances = new_page.get_text("dict", flags=fitz.TEXT_PRESERVE_IMAGES)
                                for block in text_instances.get('blocks', []):
                                    for line in block.get('lines', []):
                                        for span in line.get('spans', []):
                                            if ww['word'] in span.get('text', ''):
                                                font_size = span.get('size', 11)
                                                break
                                    
                            # Add redaction first
                            new_page.add_redact_annot(rect)
                            new_page.apply_redactions()
                            
                            # Insert the new text at the same position
                            # Calculate text width and adjust if needed
                            text_width = fitz.get_text_length(strong_text, fontname='helv', fontsize=font_size)
                            rect_width = bbox.width
                            
                            # If the new text is wider than the original, adjust font size
                            if text_width > rect_width and rect_width > 0:
                                # Reduce font size proportionally, but not below 8pt
                                new_font_size = max(8, font_size * (rect_width / text_width * 0.9))
                                font_size = new_font_size
                            
                            # Insert the text with proper alignment
                            new_page.insert_text(
                                point=bbox.tl,  # Top-left corner of the original text
                                text=strong_text,
                                fontsize=font_size,
                                fontname='helv',  # Use a standard font
                                color=(0, 0, 0),  # Black text
                                overlay=True
                            )
                    
                    new_doc.save(temp_path)
                    new_doc.close()
                    
                    # Read the modified file content
                    with open(temp_path, 'rb') as f:
                        content = f.read()
                    
                    # Create a response with the modified content
                    response = Response(
                        content,
                        mimetype='application/pdf',
                        headers={
                            'Content-Disposition': f'attachment; filename=improved_{filename}',
                            'Content-Length': len(content)
                        }
                    )
                    
                    return response
                else:
                    return jsonify({'weak_words': weak_words, 'total': len(weak_words)})
                    
        finally:
            # Clean up the temporary file copy
            if os.path.exists(temp_filepath):
                try:
                    os.unlink(temp_filepath)
                except:
                    pass
                
    except Exception as e:
        return jsonify({
            'error': 'Failed to process document',
            'details': str(e),
            'type': type(e).__name__
        }), 500
        
    finally:
        # Clean up temporary files and directories
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass
                
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass

@app.route('/highlight', methods=['POST'])
def highlight_text():
    data = request.json
    filename = data.get('filename')
    term = data.get('term')
    if not filename or not term:
        return jsonify({'error': 'Missing filename or term'}), 400
    
    BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
    OUTPUT_FOLDER = os.path.join(BASE_DIR, 'static')
    pdf_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(pdf_path):
        return jsonify({'error': 'PDF file not found.'}), 404
    
    doc = fitz.open(pdf_path)
    results = []
    for i, page in enumerate(doc):
        matches = page.search_for(term)
        if matches:
            results.append({
                "page": i + 1,
                "matches": len(matches)
            })
            for rect in matches:
                highlight = page.add_highlight_annot(rect)
                highlight.update()
    output_filename = f"highlighted_{filename}"
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)
    doc.save(output_path)
    doc.close()
    return jsonify({'message': 'Highlighting complete', 'output_filename': output_filename, 'results': results})

@app.route('/convert-and-upload', methods=['POST'])
def convert_and_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.lower().endswith('.pdf'):
        filename = secure_filename(file.filename)
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(upload_path)
        # Convert PDF to Word
        docx_filename = filename.rsplit('.', 1)[0] + '.docx'
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        try:
            with pdfplumber.open(upload_path) as pdf:
                document = Document()
                style = document.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.color.rgb = RGBColor(0, 0, 0)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        for line in page_text.split('\n'):
                            para = document.add_paragraph(line)
                            run = para.runs[0]
                            run.font.name = 'Calibri'
                            run.font.color.rgb = RGBColor(0, 0, 0)
                document.save(docx_path)
        except Exception as e:
            return jsonify({'error': f'Conversion failed: {str(e)}'}), 500
        return jsonify({'message': 'PDF converted and uploaded as Word file', 'docx_filename': docx_filename})
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/replace-in-docx', methods=['POST'])
def replace_in_docx():
    data = request.json
    filename = data.get('filename')  # PDF filename
    search_text = data.get('searchText')
    replace_text = data.get('replaceText')
    if not all([filename, search_text, replace_text]):
        return jsonify({'error': 'Missing required parameters'}), 400
    docx_filename = filename.rsplit('.', 1)[0] + '.docx'
    docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
    if not os.path.exists(docx_path):
        return jsonify({'error': 'Word file not found'}), 404
    try:
        from docx import Document
        document = Document(docx_path)
        for para in document.paragraphs:
            if search_text in para.text:
                for run in para.runs:
                    run.text = run.text.replace(search_text, replace_text)
        document.save(docx_path)
        # Convert modified docx to PDF
        new_pdf_filename = f"modified_{filename}"
        new_pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], new_pdf_filename)
        docx2pdf_convert(docx_path, new_pdf_path)
        return jsonify({'message': 'Replacement done and new PDF generated', 'pdf_filename': new_pdf_filename})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def extract_rich_text(pdf_path):
    """Extract text with formatting and structure from PDF."""
    result = {
        'pages': [],
        'metadata': {}
    }
    
    # Use context manager to ensure the document is properly closed
    with fitz.open(pdf_path) as doc:
        result['metadata'] = dict(doc.metadata)
        
        for page_num, page in enumerate(doc):
            try:
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_IMAGES)
                if not isinstance(blocks, dict) or 'blocks' not in blocks:
                    continue
                    
                blocks = blocks["blocks"]
                
                page_content = {
                    'page_number': page_num + 1,
                    'dimensions': {
                        'width': page.rect.width,
                        'height': page.rect.height
                    },
                    'content': []
                }
                
                for block in blocks:
                    if not isinstance(block, dict):
                        continue
                        
                    if 'lines' in block:
                        block_content = {
                            'type': 'text',
                            'bbox': block.get('bbox', []),
                            'lines': []
                        }
                        
                        for line in block.get('lines', []):
                            if not isinstance(line, dict):
                                continue
                                
                            line_content = {
                                'bbox': line.get('bbox', []),
                                'spans': []
                            }
                            
                            for span in line.get('spans', []):
                                if not isinstance(span, dict):
                                    continue
                                    
                                try:
                                    font = span.get('font', '').lower()
                                    is_heading = span.get('size', 0) > 11
                                    span_info = {
                                        'text': span.get('text', ''),
                                        'font': {
                                            'name': span.get('font', ''),
                                            'size': span.get('size', 0),
                                            'is_bold': 'bold' in font,
                                            'is_italic': 'italic' in font
                                        },
                                        'color': span.get('color', 0),
                                        'is_heading': is_heading
                                    }
                                    line_content['spans'].append(span_info)
                                except Exception as e:
                                    print(f"Error processing span: {e}")
                            
                            if line_content['spans']:  # Only add if we have spans
                                block_content['lines'].append(line_content)
                        
                        if block_content['lines']:  # Only add if we have lines
                            page_content['content'].append(block_content)
                    
                    elif 'image' in block:
                        try:
                            xref = block.get('image')
                            if xref is not None:
                                base_image = doc.extract_image(xref)
                                if base_image and 'image' in base_image:
                                    image_bytes = base_image['image']
                                    image_b64 = base64.b64encode(image_bytes).decode('utf-8')
                                    
                                    image_info = {
                                        'type': 'image',
                                        'bbox': block.get('bbox', []),
                                        'width': block.get('width', 0),
                                        'height': block.get('height', 0),
                                        'bpc': block.get('bpc', 8),
                                        'image': image_b64,
                                        'ext': base_image.get('ext', 'png')
                                    }
                                    page_content['content'].append(image_info)
                        except Exception as e:
                            print(f"Error extracting image: {e}")
            
            except Exception as e:
                print(f"Error processing page {page_num + 1}: {e}")
                continue
            
            if page_content['content']:  # Only add if we have content
                result['pages'].append(page_content)
    
    return result

def format_as_html(rich_text_data):
    """Convert rich text data to formatted HTML."""
    html = ['<html><head><style>']
    html.append('''
        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 0; padding: 20px; background: #f5f5f5; }
        .page {
            position: relative;
            margin: 20px auto;
            background: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            padding: 20px;
            box-sizing: border-box;
            max-width: 100%;
        }
        .heading {
            font-size: 1.4em;
            font-weight: bold;
            margin: 20px 0 10px 0;
            color: #2c3e50;
        }
        p {
            margin: 8px 0;
            line-height: 1.6;
        }
        img {
            max-width: 100%;
            height: auto;
            margin: 10px 0;
            border: 1px solid #eee;
        }
        .bold { font-weight: bold; }
        .italic { font-style: italic; }
    ''')
    html.append('</style></head><body>')
    
    for page in rich_text_data['pages']:
        html.append(f'<div class="page" style="width:{page["dimensions"]["width"]}px; max-width:100%;">')
        
        for block in page['content']:
            if block['type'] == 'text':
                for line in block['lines']:
                    line_html = []
                    for span in line['spans']:
                        styles = []
                        classes = []
                        
                        if span['is_heading']:
                            classes.append('heading')
                        if span['font']['is_bold']:
                            classes.append('bold')
                        if span['font']['is_italic']:
                            classes.append('italic')
                        if span['color'] != 0:
                            styles.append(f'color:#{span["color"]:06x}')
                                
                        style_attr = ' style="{}"'.format(';'.join(styles)) if styles else ''
                        class_attr = ' class="{}"'.format(' '.join(classes)) if classes else ''
                        
                        # Escape HTML special characters
                        text = span['text'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        line_html.append(f'<span{style_attr}{class_attr}>{text}</span>')
                    
                    if any(span['is_heading'] for span in line['spans']):
                        html.append(f'<h3>{"".join(line_html)}</h3>')
                    else:
                        html.append(f'<p>{"".join(line_html)}</p>')
            
            elif block['type'] == 'image' and 'image' in block:
                html.append(f'<img src="data:image/{block.get("ext", "png")};base64,{block["image"]}" style="max-width:100%;">')
        
        html.append('</div>')
    
    html.append('</body></html>')
    return '\n'.join(html)

@app.route('/extract-rich-text', methods=['POST'])
def extract_rich_text_route():
    """API endpoint to extract rich text from PDF."""
    temp_file_path = None
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file or file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    try:
        # Create a temporary file with a unique name
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name
            
        # Process the file
        rich_text = extract_rich_text(temp_file_path)
        
        # Get the desired output format
        format_type = request.args.get('format', 'json').lower()
        
        # Return the appropriate response format
        if format_type == 'html':
            return Response(
                format_as_html(rich_text),
                mimetype='text/html',
                headers={'Content-Disposition': 'inline'}
            )
        
        return jsonify(rich_text)
    
    except Exception as e:
        error_type = type(e).__name__
        error_msg = str(e)
        print(f"Error in extract_rich_text_route: {error_type} - {error_msg}")
        
        return jsonify({
            'error': 'Failed to process PDF',
            'details': error_msg,
            'type': error_type
        }), 500
        
    finally:
        # Clean up the temporary file if it exists
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as cleanup_error:
                print(f"Warning: Could not delete temporary file {temp_file_path}: {cleanup_error}")

if __name__ == '__main__':
    app.run(debug=True) 
